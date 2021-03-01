import PySimpleGUI as sg # импорт графической библеотеки
import pyaudio # импорт звуковой библеотеки
import numpy as np # импорт библеотеки для рассчетов частоты
from docx import Document # импорт библеотеки для вывода результатов работы программы

def error_win(): # Функция ошибок - содержит варианты всех неправильных действий пользователя.
    sg.theme('DarkBlue')
    layout = [
        [sg.Text(' ')],
        [sg.Text(' ' * 20), sg.Text('ОШИБКА', text_color='red')],
        [sg.Text(' ')],
        [sg.Text('Возможно вы: ')],
        [sg.Text('1. Не ввели длительность вашей записи или ввели недопустимое значение')],
        [sg.Text('2. Не подключили записывающее устройство')],
        [sg.Text(' ')],
        [sg.Text(' ' * 10), sg.Text('Попробуйте устранить неполадки и перезапустите программу')],
        [sg.Text(' ')],
    ]
    window = sg.Window('ERROR', layout).Finalize()
    while True:
        event, values = window.read()
        if event in (None, 'Exit'):
            break


def freq_to_note(frequency): # Функция сопоставляет значение частоты звука с нотой

    midi_number = (69 + 12*np.log2(frequency/440.0)) # формула выведена на основе данных сайта: 
    n = midi_number.astype(np.int)                      #https://newt.phys.unsw.edu.au/jw/notes.html
    notes = 'C C# D D# E F F# G G# A A# B'.split()                                                    
    return notes[n % 12] # + str(round(n/12 - 1)) <- следует добавить в return при необходимости определения октавы.
    
def notation_to_out(arr): # Функция, получая на ввод массив нот, создает word-файл,
                         # вставляет в него последовательно картинки нот, 
                        #сопоставляя их с каждым элементом массива.
    document = Document()
    
    p = document.add_paragraph()
    r = p.add_run()

    r.add_picture('pics/sk.png') #Скрипичный ключ
    
    i=1
    while i<len(arr):
        if arr[i] == 'A' : r.add_picture('pics/a.png')
        elif arr[i] == 'A#' : r.add_picture('pics/a1.png')
        elif arr[i] == 'B' : r.add_picture('pics/b.png')
        elif arr[i] == 'C' : r.add_picture('pics/c.png')
        elif arr[i] == 'C#' : r.add_picture('pics/c1.png')
        elif arr[i] == 'D' : r.add_picture('pics/d.png')
        elif arr[i] == 'D#' : r.add_picture('pics/d1.png')
        elif arr[i] == 'E' : r.add_picture('pics/e.png')
        elif arr[i] == 'F' : r.add_picture('pics/f.png')
        elif arr[i] == 'F#' : r.add_picture('pics/f1.png')
        elif arr[i] == 'G' : r.add_picture('pics/g.png')
        elif arr[i] == 'G#' : r.add_picture('pics/g1.png')
        i+=1

    document.save('Результат программы.docx')

def work(duration): # главная функция программы - записывая звук с микрофона,
                    # определяет частоту звука, составляет массив из нот, обработанных функцией
    chunk = 4096 
    fs = 44100

    A = [] # создание массива для нот

    p = pyaudio.PyAudio()
    try:
        stream=p.open(format=pyaudio.paInt16,
                      channels=1,
                      rate=fs,
                      input=True,
                      frames_per_buffer=chunk)
    
    except IOError as e: error_win() 
    
    for i in range(0, int(fs / chunk * duration)):  # запись и нахождение частоты 
                                                    #при помощи Быстрых Преобразований Фурье - далее БПФ
        indata = np.fromstring(stream.read(chunk),dtype=np.int16)
 
        # Применение БПФ
        fftData=abs(np.fft.rfft(indata))**2
        # Нахождение максимума
        which = fftData[1:].argmax() + 1
        # Использование квадратичной интерполяции в окресности точки максимума
        if which != len(fftData)-1:
            y0,y1,y2 = np.log(fftData[which-1:which+2:])
            x1 = (y2 - y0) * .5 / (2 * y1 - y2 - y0)
        # Нахождение частоты, вызов freq_to_note, добавление элементов в массив
            frequency = (which+x1)*fs/chunk
            note = freq_to_note(frequency)
            A.append(note)
        else:
            frequency = which*fs/chunk
            note = freq_to_note(frequency)
            A.append(note)
     
    stream.close()
    p.terminate()

    i=1  # т.к. частота дискретизации большая, то мы получаем 
    while i<len(A): # множество повторяющихся элементов массива подряд
        if A[i-1] == A[i]: del A[i] # удалим их
        i+=1

    notation_to_out(A)

sg.theme('DarkBlue') # Начальное окно программы
layout = [
    [sg.Text(' ')],
    [sg.Text('Преобразователь звуковой дорожки в нотную запись')],
    [sg.Text(' ')],    
    [sg.Text('Данная программа выполняет преобразование вашего проигранного звука в нотный текст')],
    [sg.Text('Для начала работы с программой вам потребуется лишь записывающее устройство, убедитесь, что оно подключено')],
    [sg.Text('Принцип работы программы очень прост - вам всего лишь необходимо ввести во всплывающее окно примерную длительность вашей записи')],
    [sg.Text('далее - начнется автоматическая запись, а после программа сохранит результ работы в свою директорию в виде word-файла')],
    [sg.Text(' ')],
    [sg.Text(' ' * 90), sg.Button('Начать!', size=(10,3))],
    [sg.Text(' ')],
]

window = sg.Window(' ', layout).Finalize()
 
while True:  
    event, values = window.read()
    if event in (None, 'Exit'):
        break
    
    if  event == 'Начать!' :  #Действия при нажатии кнопки "Начать" 
        window.close()
        try:
            duration = int(sg.PopupGetText('Введите примерную длительность вашей записи (в секундах): '))
        except TypeError as e: error_win()
        if duration <= 0: error_win()
        work(duration)
